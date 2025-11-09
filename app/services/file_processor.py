"""
File processing service for text extraction
"""

import io
import re
from typing import Optional, Dict, Any
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import logging

logger = logging.getLogger(__name__)

class FileProcessor:
    """
    Handles text extraction from various file formats
    """
    
    def extract_text(self, file_content: bytes, mime_type: str) -> Optional[str]:
        """
        Extract text from file content based on MIME type
        """
        if mime_type == 'application/pdf':
            return self._extract_pdf_text(file_content)
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return self._extract_docx_text(file_content)
        elif mime_type in ['text/plain', 'text/markdown']:
            return self._extract_text_file(file_content)
        else:
            return None
    
    def extract_text_with_pages(self, file_content: bytes, mime_type: str) -> Optional[Dict[str, Any]]:
        """
        Extract text with page information for better chunking
        """
        if mime_type == 'application/pdf':
            return self._extract_pdf_text_with_pages(file_content)
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return self._extract_docx_text_with_pages(file_content)
        elif mime_type in ['text/plain', 'text/markdown']:
            return self._extract_text_file_with_pages(file_content)
        else:
            return None
    
    def _extract_pdf_text(self, pdf_content: bytes) -> Optional[str]:
        """Extract text from PDF"""
        try:
            pdf_reader = PdfReader(io.BytesIO(pdf_content))
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    # Clean up spacing issues
                    cleaned_text = self._clean_text_spacing(text)
                    text_parts.append(cleaned_text)
            
            return '\n\n'.join(text_parts) if text_parts else None
            
        except Exception:
            return None
    
    def _clean_text_spacing(self, text: str) -> str:
        """Clean up common PDF text extraction spacing issues"""
        # Fix multiple spaces
        text = re.sub(r' +', ' ', text)
        # Fix spaces before punctuation
        text = re.sub(r' +([.!?,:;])', r'\1', text)
        # Fix line breaks in the middle of words
        text = re.sub(r'(\w)-\s*\n\s*(\w)', r'\1\2', text)
        # Fix multiple line breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        return text.strip()
    
    def _extract_docx_text(self, docx_content: bytes) -> Optional[str]:
        """Extract text from DOCX"""
        try:
            doc = DocxDocument(io.BytesIO(docx_content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            return '\n\n'.join(text_parts) if text_parts else None
            
        except Exception:
            return None
    
    def _extract_text_file(self, text_content: bytes) -> Optional[str]:
        """Extract text from plain text or markdown"""
        try:
            # Try UTF-8 first, fallback to latin-1
            try:
                return text_content.decode('utf-8')
            except UnicodeDecodeError:
                return text_content.decode('latin-1')
        except Exception:
            return None
    
    def _extract_pdf_text_with_pages(self, pdf_content: bytes) -> Optional[Dict[str, Any]]:
        """Extract text from PDF with page information, including tables"""
        try:
            # Try pdfplumber first for better table extraction
            try:
                import pdfplumber
                return self._extract_pdf_with_pdfplumber(pdf_content)
            except ImportError:
                logger.warning("pdfplumber not available, falling back to PyPDF2")
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}, falling back to PyPDF2")
            
            # Fallback to PyPDF2
            pdf_reader = PdfReader(io.BytesIO(pdf_content))
            pages = []
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    cleaned_text = self._clean_text_spacing(text)
                    pages.append({
                        'page_number': page_num,
                        'text': cleaned_text
                    })
            
            if pages:
                full_text = '\n\n'.join([page['text'] for page in pages])
                return {
                    'text': full_text,
                    'pages': pages,
                    'total_pages': len(pages)
                }
            return None
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return None
    
    def _extract_pdf_with_pdfplumber(self, pdf_content: bytes) -> Optional[Dict[str, Any]]:
        """Extract text from PDF using pdfplumber (better table support)"""
        import pdfplumber
        
        pages = []
        pdf = pdfplumber.open(io.BytesIO(pdf_content))
        
        for page_num, page in enumerate(pdf.pages, 1):
            page_text_parts = []
            
            # Extract regular text
            text = page.extract_text()
            if text and text.strip():
                cleaned_text = self._clean_text_spacing(text)
                page_text_parts.append(cleaned_text)
            
            # Extract tables and convert to text format
            tables = page.extract_tables()
            for table in tables:
                if table:
                    # Convert table to markdown-like format with pipes
                    table_lines = []
                    for row in table:
                        if row and any(cell for cell in row if cell):
                            # Filter out None values and clean cells
                            cleaned_row = [str(cell).strip() if cell else '' for cell in row]
                            if any(cleaned_row):  # Only add non-empty rows
                                table_line = ' | '.join(cleaned_row)
                                table_lines.append(f'| {table_line} |')
                    
                    if table_lines:
                        page_text_parts.append('\n'.join(table_lines))
            
            if page_text_parts:
                page_text = '\n\n'.join(page_text_parts)
                pages.append({
                    'page_number': page_num,
                    'text': page_text
                })
        
        pdf.close()
        
        if pages:
            full_text = '\n\n'.join([page['text'] for page in pages])
            return {
                'text': full_text,
                'pages': pages,
                'total_pages': len(pages)
            }
        return None
    
    def _extract_docx_text_with_pages(self, docx_content: bytes) -> Optional[Dict[str, Any]]:
        """Extract text from DOCX with page information (simplified)"""
        try:
            doc = DocxDocument(io.BytesIO(docx_content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            if text_parts:
                full_text = '\n\n'.join(text_parts)
                return {
                    'text': full_text,
                    'pages': [{'page_number': 1, 'text': full_text}],
                    'total_pages': 1
                }
            return None
            
        except Exception:
            return None
    
    def _extract_text_file_with_pages(self, text_content: bytes) -> Optional[Dict[str, Any]]:
        """Extract text from plain text or markdown with page information"""
        try:
            # Try UTF-8 first, fallback to latin-1
            try:
                text = text_content.decode('utf-8')
            except UnicodeDecodeError:
                text = text_content.decode('latin-1')
            
            if text:
                return {
                    'text': text,
                    'pages': [{'page_number': 1, 'text': text}],
                    'total_pages': 1
                }
            return None
            
        except Exception:
            return None
